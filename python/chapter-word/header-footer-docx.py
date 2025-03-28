from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = Document()

# Add Header
doc.add_heading('页眉页脚', 0)


# add paragraph
paragraph = doc.add_paragraph()
paragraph.add_run('this is paragraph 1.')


# set header
header = doc.sections[0].header
print(len(header.paragraphs))

paragraph = header.paragraphs[0]
paragraph.add_run('这是第一页的页眉')

footer = doc.sections[0].footer
paragraph = footer.paragraphs[0]
paragraph.add_run('这是第一页的页脚')


# 页面和页脚会显示了“与上一节相同”。
# 如果不使用上一节的内容和样式要将header.is_linked_to_previous的属性或footer.is_linked_to_previous的属性设置为False，用于解除“链接上一节页眉”或者“链接上一节页脚”。
# add two sections
doc.add_section()
doc.add_section()

# 设置第二页页眉页脚
header = doc.sections[1].header
header.is_linked_to_previous = False
paragraph = header.paragraphs[0]
paragraph.add_run('这是第二页的页眉')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

footer = doc.sections[1].footer
footer.is_linked_to_previous = False
paragraph = footer.paragraphs[0]
paragraph.add_run('这是第二页的页脚')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Save the document
doc.save('example.docx')
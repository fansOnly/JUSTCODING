from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


# 字体函数
def set_font(doc, text, font_cn, font_en = None):
  style_add = doc.styles.add_style(font_cn, WD_STYLE_TYPE.CHARACTER)
  style_add.font.name = font_cn
  doc.styles[font_cn]._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
  _paragraph = doc.add_paragraph()
  _run = _paragraph.add_run(text, style=font_cn)
  if font_en:
    font = _run.font
    font.name = font_en

doc = Document()


# set font globally
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# set font for run
# 设置字体样式 - 宋体
style_font = doc.styles.add_style('宋体', WD_STYLE_TYPE.CHARACTER)
style_font.font.name = '宋体'
doc.styles['宋体']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 设置字体样式 - 楷体
style_font = doc.styles.add_style('楷体', WD_STYLE_TYPE.CHARACTER)
style_font.font.name = '楷体'
doc.styles['楷体']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

# 设置字体样式 - 黑体
style_font = doc.styles.add_style('黑体', WD_STYLE_TYPE.CHARACTER)
style_font.font.name = '黑体'
doc.styles['黑体']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')



# add title
# 添加标题（0相当于文章的题目，默认级别是1，级别范围为0-9，0时候自动带下划线）
header_1 =doc.add_heading('设置字体', 0)


# add paragraph
paragraph1 = doc.add_paragraph()

# set en font for paragraph
# 设置段落英文字体为Arial
run = paragraph1.add_run('这里全部是宋体，abcdefg\n', style='宋体')
run1 = paragraph1.add_run('这里中文是宋体，abcdefg，英文字体为Arial\n', style='宋体')
font1 = run1.font
font1.name = 'Arial'

run3 = paragraph1.add_run('这里是楷体\n', style='楷体')
run4 = paragraph1.add_run('这里是黑体\n', style='黑体')

# set font by set_font
a = '小朋友 你是否有很多问号\n'
b = '我是谁?\nHow do you do?\n'

set_font(doc, a, '仿宋')
set_font(doc, b, '微软雅黑', 'Cambira')


# save document
doc.save('example.docx')
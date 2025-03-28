from docx import Document
from docx.shared import Inches

document = Document()

# Add a heading
document.add_heading('插入图片和表格', 0)

# add image
document.add_picture('2.jpg', width=Inches(2.0), height=Inches(1.0))


# add table
table = document.add_table(rows=2,cols=1)
table.style = 'Light Grid Accent 1'

# set table values
table.cell(0,0).text = '姓名'
table.rows[1].cells[0].text = '张三'

# add table row
add_cells = table.add_row().cells
add_cells[0].text = '年龄'


document.save('example.docx')
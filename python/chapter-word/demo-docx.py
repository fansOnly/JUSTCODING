from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Inches, RGBColor, Length

doc = Document()

# set font globally
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# add title
# 添加标题（0相当于文章的题目，默认级别是1，级别范围为0-9，0时候自动带下划线）
header_1 =doc.add_heading('Document Title', 0)
doc.add_heading('Sub-Title', 1)
doc.add_heading('Sub-Sub-Title', 2)
# set heading alignment
header_1.alignment = WD_ALIGN_PARAGRAPH.CENTER


# add paragraph
# 设置段落格式：首行缩进0.75cm，居左，段后距离1.0英寸,1.5倍行距。
paragraph1 = doc.add_paragraph()
paragraph1.paragraph_format.first_line_indent = Cm(0.75)
paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph1.paragraph_format.space_after = Inches(1.0)
paragraph1.paragraph_format.line_spacing = 1.5

text1 = '中国台湾华语流行歌手、' \
       '音乐创作家、作曲家、作词人、' \
       '制作人、杰威尔音乐公司老板之一、导演。' \
       '近年涉足电影行业。周杰伦是2000年后亚洲流行乐坛最具革命性与指标' \
       '性的创作歌手，有“亚洲流行天王”之称。他突破原有亚洲音乐的主题、形' \
       '式，融合多元的音乐素材，创造出多变的歌曲风格，尤以融合中西式曲风的嘻哈' \
       '或节奏蓝调最为著名，可说是开创华语流行音乐“中国风”的先声。周杰伦的' \
       '出现打破了亚洲流行乐坛长年停滞不前的局面，为亚洲流行乐坛翻开了新的一页！\n'

run1 = paragraph1.add_run(text1)
run1.font.size = Pt(10)  # 设置字体大小
run1.font.color.rgb = RGBColor(255,0,0)  # 设置字体颜色
run1.bold = True  # 设置字体加粗
run1.italic = True  # 设置字体斜体

# print(len(paragraph1.runs))  # 输出段落中字数
# print(paragraph1.runs[0].text)  # 输出段落中第一个字的文本
# print(paragraph1.runs[-1].text)  # 输出段落中最后一个字的文本


# paragraph 2
paragraph2 = doc.add_paragraph('This is paragraph 2.')
paragraph2.add_run('bold text').bold = True
paragraph2.add_run('italic text').italic = True


# insert paragraph
paragraph3 = paragraph2.insert_paragraph_before('This is a paragraph 3 insert before paragraph 2.')


# insert page break
doc.add_page_break()


# paragraph 3
paragraph3 = doc.add_paragraph('This is a new paragraph 4 in page 2.')


# add other styles
# 无序列表
doc.add_paragraph('无序列表：', style='List Bullet')
# 有序列表
doc.add_paragraph('有序列表：', style='List Number')



# 设置字体属性
#all_caps:全部大写字母
#bold:加粗
#color:字体颜色

#double_strike:双删除线
#hidden : 隐藏
#imprint : 印记
#italic : 斜体
#name  :字体
#shadow  :阴影
#strike  :  删除线
#subscript  :下标	
#superscript  :上标
#underline  :下划线

# 对齐设置 = WD_ALIGN_PARAGRAPH
#LEFT: 左对齐
#CENTER: 文字居中
#RIGHT: 右对齐
#JUSTIFY: 文本两端对齐

# 设置段落行距 = Length
# SINGLE :单倍行距（默认）
#ONE_POINT_FIVE : 1.5倍行距
# DOUBLE2 : 倍行距
#AT_LEAST : 最小值
#EXACTLY:固定值
# MULTIPLE : 多倍行距

# save document
doc.save('example.docx')